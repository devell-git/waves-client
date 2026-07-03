#!/usr/bin/env python3
"""HTML com inline styles → DOCX formatado.

Lê HTML com styles computados (inline) e gera DOCX preservando:
- Cores de texto e fundo
- Bold, italic, font-size
- Tabelas com cores de header e células
- Listas bullet/numbered
- Headings
- Badges (inline-block com background colorido)
"""
import re
import sys
from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH


def parse_color(css_color: str) -> RGBColor | None:
    """Converte css color (rgb/rgba/hex) → RGBColor."""
    if not css_color:
        return None
    css_color = css_color.strip()
    # rgb(r, g, b) or rgba(r, g, b, a)
    m = re.match(r'rgba?\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)', css_color)
    if m:
        r, g, b = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if r == 0 and g == 0 and b == 0:
            return None  # skip black (default)
        return RGBColor(r, g, b)
    # #hex
    m = re.match(r'#([0-9a-fA-F]{6})', css_color)
    if m:
        h = m.group(1)
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))
    m = re.match(r'#([0-9a-fA-F]{3})$', css_color)
    if m:
        h = m.group(1)
        return RGBColor(int(h[0]*2, 16), int(h[1]*2, 16), int(h[2]*2, 16))
    return None


def get_style(el, prop: str) -> str:
    """Extrai propriedade do style inline."""
    style = el.get('style', '') if hasattr(el, 'get') else ''
    for part in style.split(';'):
        part = part.strip()
        if ':' in part:
            k, v = part.split(':', 1)
            if k.strip() == prop:
                return v.strip()
    return ''


def get_font_weight(el) -> bool:
    """Retorna True se bold."""
    fw = get_style(el, 'font-weight')
    return fw in ('bold', '700', '800', '900', '600')


def get_font_size(el) -> float | None:
    """Retorna tamanho em pt."""
    fs = get_style(el, 'font-size')
    m = re.match(r'([\d.]+)px', fs)
    if m:
        return float(m.group(1)) * 0.75  # px to pt
    m = re.match(r'([\d.]+)pt', fs)
    if m:
        return float(m.group(1))
    return None


def set_cell_bg(cell, color_str: str):
    """Define cor de fundo de uma célula da tabela."""
    color = parse_color(color_str)
    if not color:
        return
    hex_color = f"{color[0]:02X}{color[1]:02X}{color[2]:02X}"
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_run_with_style(paragraph, text: str, el=None):
    """Adiciona run com formatação baseada nos inline styles."""
    if not text.strip():
        return
    run = paragraph.add_run(text)

    if el is None:
        return run

    # Bold
    if get_font_weight(el):
        run.bold = True
    tag = getattr(el, 'name', '')
    if tag in ('strong', 'b'):
        run.bold = True
    if tag in ('em', 'i'):
        run.italic = True

    # Font size
    fs = get_font_size(el)
    if fs and fs > 5:
        run.font.size = Pt(fs)

    # Text color
    color = parse_color(get_style(el, 'color'))
    if color:
        run.font.color.rgb = color

    # Code font
    if tag == 'code' or 'monospace' in get_style(el, 'font-family'):
        run.font.name = 'Consolas'
        if not fs:
            run.font.size = Pt(9)

    # Background color as highlight (badges)
    bg = get_style(el, 'background-color') or get_style(el, 'background')
    if bg and 'rgba(0' not in bg and 'transparent' not in bg:
        bg_color = parse_color(bg)
        if bg_color:
            # Add text with background - use shading on run
            rpr = run._r.get_or_add_rPr()
            hex_bg = f"{bg_color[0]:02X}{bg_color[1]:02X}{bg_color[2]:02X}"
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_bg}" w:val="clear"/>')
            rpr.append(shading)

    return run


def add_formatted_text(paragraph, element):
    """Recursivamente adiciona texto formatado."""
    for child in element.children:
        if isinstance(child, str):
            text = child
            if text.strip():
                add_run_with_style(paragraph, text, element)
        elif child.name in ('strong', 'b'):
            run = paragraph.add_run(child.get_text())
            run.bold = True
            color = parse_color(get_style(child, 'color'))
            if color:
                run.font.color.rgb = color
        elif child.name in ('em', 'i'):
            run = paragraph.add_run(child.get_text())
            run.italic = True
        elif child.name == 'code':
            add_run_with_style(paragraph, child.get_text(), child)
        elif child.name == 'br':
            paragraph.add_run('\n')
        elif child.name == 'span':
            add_run_with_style(paragraph, child.get_text(), child)
        elif child.name in ('div', 'p'):
            # Inline div/p within a paragraph — just add text
            text = child.get_text().strip()
            if text:
                add_run_with_style(paragraph, text, child)
                paragraph.add_run(' ')
        else:
            text = child.get_text().strip()
            if text and child.name not in ('script', 'style', 'svg', 'path', 'button'):
                add_run_with_style(paragraph, text, child)


def process_element(doc, el):
    """Processa um elemento HTML em docx."""
    if isinstance(el, str):
        text = el.strip()
        if text:
            doc.add_paragraph(text)
        return

    tag = getattr(el, 'name', None)
    if not tag or tag in ('script', 'style', 'svg', 'path', 'button', 'input'):
        return

    # Headings
    if tag in ('h1', 'h2', 'h3', 'h4'):
        level = int(tag[1])
        heading = doc.add_heading('', level=level)
        add_formatted_text(heading, el)
        return

    # Paragraph
    if tag == 'p':
        p = doc.add_paragraph()
        add_formatted_text(p, el)
        # Background on paragraph
        bg = get_style(el, 'background-color') or get_style(el, 'background')
        if bg and 'transparent' not in bg and 'rgba(0' not in bg:
            bg_color = parse_color(bg)
            if bg_color:
                hex_bg = f"{bg_color[0]:02X}{bg_color[1]:02X}{bg_color[2]:02X}"
                ppr = p._p.get_or_add_pPr()
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_bg}" w:val="clear"/>')
                ppr.append(shading)
        return

    # Lists
    if tag == 'ul':
        for li in el.find_all('li', recursive=False):
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, li)
        return

    if tag == 'ol':
        for li in el.find_all('li', recursive=False):
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, li)
        return

    # Tables
    if tag == 'table':
        rows = el.find_all('tr')
        if not rows:
            return
        cols = rows[0].find_all(['th', 'td'])
        if not cols:
            return
        table = doc.add_table(rows=0, cols=len(cols))
        table.style = 'Table Grid'

        for tr in rows:
            cells = tr.find_all(['th', 'td'])
            row = table.add_row()
            for i, cell in enumerate(cells):
                if i >= len(row.cells):
                    break
                # Cell text
                row.cells[i].text = ''
                p = row.cells[i].paragraphs[0]
                add_formatted_text(p, cell)

                # Header styling
                if cell.name == 'th':
                    for run in p.runs:
                        run.bold = True
                    set_cell_bg(row.cells[i], '#f1f5f9')

                # Cell background from inline style
                cell_bg = get_style(cell, 'background-color') or get_style(cell, 'background')
                if cell_bg and 'transparent' not in cell_bg:
                    set_cell_bg(row.cells[i], cell_bg)
        return

    # Blockquote
    if tag == 'blockquote':
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        add_formatted_text(p, el)
        for run in p.runs:
            run.italic = True
        # Left border via background
        bg = get_style(el, 'background-color') or get_style(el, 'background')
        if bg:
            bg_color = parse_color(bg)
            if bg_color:
                hex_bg = f"{bg_color[0]:02X}{bg_color[1]:02X}{bg_color[2]:02X}"
                ppr = p._p.get_or_add_pPr()
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_bg}" w:val="clear"/>')
                ppr.append(shading)
        return

    # Pre/code block
    if tag == 'pre':
        p = doc.add_paragraph()
        run = p.add_run(el.get_text())
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        set_cell_bg_para = parse_color('#f8fafc')
        if set_cell_bg_para:
            ppr = p._p.get_or_add_pPr()
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8FAFC" w:val="clear"/>')
            ppr.append(shading)
        return

    # HR
    if tag == 'hr':
        p = doc.add_paragraph()
        p.add_run('─' * 60).font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
        return

    # Details (colapsável — forçar aberto)
    if tag == 'details':
        summary = el.find('summary')
        if summary:
            p = doc.add_paragraph()
            run = p.add_run(summary.get_text().strip())
            run.bold = True
        for child in el.children:
            if getattr(child, 'name', None) != 'summary':
                process_element(doc, child)
        return

    # Container divs — process children, apply background if present
    if tag in ('div', 'section', 'article', 'main', 'body', 'html', 'span'):
        # Check if this is a "card" or "KPI" div with background
        bg = get_style(el, 'background-color') or get_style(el, 'background')
        has_border = bool(get_style(el, 'border') or get_style(el, 'border-radius'))
        text_content = el.get_text().strip()

        # Simple text div (KPI number, label, etc.)
        children = list(el.children)
        has_block_children = any(
            getattr(c, 'name', None) in ('div', 'p', 'h1', 'h2', 'h3', 'table', 'ul', 'ol', 'details', 'pre', 'blockquote')
            for c in children
        )

        if not has_block_children and text_content and len(text_content) < 200:
            # Leaf div — render as paragraph with styles
            p = doc.add_paragraph()
            add_formatted_text(p, el)
            if bg and 'transparent' not in bg and 'rgba(0' not in bg:
                bg_color = parse_color(bg)
                if bg_color:
                    hex_bg = f"{bg_color[0]:02X}{bg_color[1]:02X}{bg_color[2]:02X}"
                    ppr = p._p.get_or_add_pPr()
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_bg}" w:val="clear"/>')
                    ppr.append(shading)
        else:
            # Container — recurse
            for child in children:
                process_element(doc, child)
        return

    # Fallback — just text
    text = el.get_text().strip()
    if text:
        p = doc.add_paragraph()
        add_formatted_text(p, el)


def html_to_docx(html_path: str, docx_path: str):
    """Converte HTML com inline styles para DOCX."""
    html = Path(html_path).read_text(encoding='utf-8')
    soup = BeautifulSoup(html, 'html.parser')

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Process body or root
    body = soup.find('body') or soup
    for child in body.children:
        process_element(doc, child)

    doc.save(docx_path)


if __name__ == '__main__':
    if len(sys.argv) != 3:
        print(f"Usage: {sys.argv[0]} input.html output.docx")
        sys.exit(1)
    html_to_docx(sys.argv[1], sys.argv[2])
